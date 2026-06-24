"""Word document helpers for coursework reports."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

A4_WIDTH_INCHES = 8.27
A4_HEIGHT_INCHES = 11.69
SIDE_MARGIN_INCHES = 1.0
VERTICAL_MARGIN_INCHES = 0.85


@dataclass(frozen=True)
class WordParagraph:
    """One extracted Word paragraph."""

    index: int
    text: str
    style: str


def _document_classes():
    """Return python-docx classes or raise a clear dependency error."""

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover - checked during setup
        raise RuntimeError("python-docx is required for Word report support") from exc
    return Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt


def _set_base_styles(document: object) -> None:
    """Apply conservative built-in Word style defaults."""

    _, _, _, _, _, pt = _document_classes()
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = pt(11)
    for style_name, size in [
        ("Title", 18),
        ("Heading 1", 15),
        ("Heading 2", 13),
        ("Heading 3", 12),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = pt(size)
        style.font.bold = True


def _configure_a4(document: object) -> None:
    """Configure all Word sections for A4 portrait coursework reports."""

    _, _, _, _, inches, _ = _document_classes()
    for section in document.sections:
        section.page_width = inches(A4_WIDTH_INCHES)
        section.page_height = inches(A4_HEIGHT_INCHES)
        section.left_margin = inches(SIDE_MARGIN_INCHES)
        section.right_margin = inches(SIDE_MARGIN_INCHES)
        section.top_margin = inches(VERTICAL_MARGIN_INCHES)
        section.bottom_margin = inches(VERTICAL_MARGIN_INCHES)


def _add_field(paragraph: object, instruction: str, placeholder: str) -> None:
    """Add a Word field, such as a TOC field, to a paragraph."""

    _, _, element, qn, _, _ = _document_classes()

    begin = element("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph._p.append(begin)

    instr = element("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = instruction
    paragraph._p.append(instr)

    separate = element("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph._p.append(separate)

    paragraph.add_run(placeholder)

    end = element("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(end)


def _placeholder(topic: str, fallback: str) -> str:
    """Return a topic-aware placeholder."""

    clean_topic = " ".join(topic.split())
    if not clean_topic:
        return fallback
    return fallback.replace("this topic", clean_topic)


def create_word_report(
    path: str | Path,
    *,
    title: str,
    authors: str = "",
    topic: str = "",
) -> Path:
    """Create a Word-first coursework report scaffold."""

    document_class, align, _, _, _, pt = _document_classes()
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = document_class()
    _configure_a4(document)
    _set_base_styles(document)
    document.core_properties.title = title
    if authors:
        document.core_properties.author = authors
    if topic:
        document.core_properties.subject = topic

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = align.CENTER
    title_paragraph.add_run(title)
    if authors:
        author_paragraph = document.add_paragraph()
        author_paragraph.alignment = align.CENTER
        author_paragraph.add_run(authors)
    course_paragraph = document.add_paragraph()
    course_paragraph.alignment = align.CENTER
    course_paragraph.add_run("FINS2026 - Fintech")

    document.add_paragraph()
    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        _placeholder(
            topic,
            "[REMOVE] Summarise this topic. State the research question, data or "
            "method, and main result in 100 to 200 words.",
        )
    )

    document.add_page_break()
    document.add_heading("Table of Contents", level=1)
    toc_paragraph = document.add_paragraph()
    _add_field(toc_paragraph, r'TOC \o "1-3" \h \z \u', "Update this table in Word.")

    sections = [
        (
            "Introduction",
            _placeholder(
                topic,
                "[REMOVE] Open with a concrete finding or research question about this topic.",
            ),
        ),
        (
            "Data",
            "[REMOVE] Describe the data source, sample period, variables, units, and filters.",
        ),
        (
            "Methodology",
            "[REMOVE] Explain the empirical design, assumptions, and model inputs.",
        ),
        (
            "Results",
            "[REMOVE] Report the main quantitative findings with tables and figures.",
        ),
        (
            "Conclusion",
            "[REMOVE] State the answer, limitations, and practical implications.",
        ),
        (
            "References",
            "[REMOVE] Insert a Word bibliography after adding sources with the References tab.",
        ),
        (
            "Appendix",
            "[REMOVE] Add supplementary robustness checks, data notes, or extra tables.",
        ),
    ]

    for heading, body in sections:
        document.add_page_break()
        document.add_heading(heading, level=1)
        paragraph = document.add_paragraph(body)
        paragraph.paragraph_format.space_after = pt(8)

    document.save(output_path)
    return output_path


def read_word_paragraphs(path: str | Path) -> list[WordParagraph]:
    """Extract non-empty paragraphs and their Word style names."""

    document_class, *_ = _document_classes()
    document = document_class(str(path))
    paragraphs: list[WordParagraph] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        paragraphs.append(WordParagraph(index=index, text=text, style=style_name))
    return paragraphs


def word_heading_level(style_name: str) -> int | None:
    """Return a heading level for built-in Word heading styles."""

    normalized = style_name.strip().lower()
    if not normalized.startswith("heading "):
        return None
    _, _, suffix = normalized.partition(" ")
    try:
        level = int(suffix)
    except ValueError:
        return None
    return level if 1 <= level <= 9 else None
