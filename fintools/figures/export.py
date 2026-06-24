"""Figure export helpers for LaTeX, Word, and student reports."""

from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt

A4_PORTRAIT_WIDTH_INCHES = 8.27
A4_PORTRAIT_HEIGHT_INCHES = 11.69
A4_SIDE_MARGIN_INCHES = 1.0
A4_VERTICAL_MARGIN_INCHES = 0.75
EMU_PER_INCH = 914400


@dataclass(frozen=True)
class FigureContext:
    """Caption and source context saved alongside generated figures."""

    title: str
    note: str = ""
    source: str = ""
    sample: str = ""
    units: str = ""

    def caption_text(self, figure_number: int | None = None) -> str:
        """Return a self-contained Word/report caption."""

        prefix = "Figure" if figure_number is None else f"Figure {figure_number}"
        parts = [f"{prefix}. {_sentence(self.title)}"]
        if self.note:
            parts.append(_sentence(self.note))
        if self.sample:
            parts.append(_sample_sentence(self.sample))
        if self.units:
            parts.append(f"Units: {_sentence(self.units)}")
        if self.source:
            parts.append(f"Source: {_sentence(self.source)}")
        return " ".join(parts)

    def as_markdown(self) -> str:
        lines = [f"# {self.title}", ""]
        if self.note:
            lines.extend(["## Note", self.note, ""])
        if self.sample:
            lines.extend(["## Sample", self.sample, ""])
        if self.units:
            lines.extend(["## Units", self.units, ""])
        if self.source:
            lines.extend(["## Source", self.source, ""])
        return "\n".join(lines).rstrip() + "\n"


@dataclass(frozen=True)
class WordFigureSpec:
    """Word/A4 export dimensions."""

    name: str
    width_inches: float
    height_inches: float | None = None
    dpi: int = 300


@dataclass(frozen=True)
class WordFigureEntry:
    """One figure entry for a combined Word proof document."""

    image_path: str | Path
    context: FigureContext | None = None
    spec: str | WordFigureSpec = "full_width"


WORD_FIGURE_SPECS = {
    "full_width": WordFigureSpec("full_width", width_inches=6.27, height_inches=3.75),
    "half_width": WordFigureSpec("half_width", width_inches=3.05, height_inches=2.25),
    "portrait_tall": WordFigureSpec("portrait_tall", width_inches=6.27, height_inches=5.05),
    "portrait_full": WordFigureSpec("portrait_full", width_inches=6.27, height_inches=7.35),
    "two_panel": WordFigureSpec("two_panel", width_inches=6.27, height_inches=4.20),
    "landscape_wide": WordFigureSpec("landscape_wide", width_inches=9.70, height_inches=5.45),
}


def _sentence(text: str) -> str:
    """Return stripped text with sentence-ending punctuation."""

    value = " ".join(str(text).strip().split())
    if not value:
        return ""
    return value if value.endswith((".", "?", "!")) else f"{value}."


def _sample_sentence(sample: str) -> str:
    """Return the standard sample-period caption sentence."""

    value = " ".join(str(sample).strip().split())
    if not value:
        return ""
    if value.lower().startswith("the sample spans"):
        return _sentence(value)
    if " to " in value:
        start, end = value.split(" to ", 1)
        return f"The sample spans the time period {start.strip()} to {end.strip()}."
    return f"Sample: {_sentence(value)}"


def word_figure_spec(name: str = "full_width") -> WordFigureSpec:
    """Return a named Word/A4 figure export spec."""

    try:
        return WORD_FIGURE_SPECS[name]
    except KeyError as exc:
        names = ", ".join(sorted(WORD_FIGURE_SPECS))
        raise ValueError(f"unknown Word figure spec '{name}'. Available: {names}") from exc


def save_figure(
    fig: plt.Figure,
    path: str | Path,
    *,
    dpi: int = 300,
    transparent: bool = False,
) -> Path:
    """Save one figure file with consistent report defaults."""

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(
        output_path,
        dpi=dpi,
        bbox_inches="tight",
        pad_inches=0.04,
        facecolor="white",
        transparent=transparent,
        metadata={"Creator": "fintools.figures"},
    )
    return output_path


def export_figure_bundle(
    fig: plt.Figure,
    output_dir: str | Path,
    stem: str,
    *,
    context: FigureContext | None = None,
    formats: tuple[str, ...] = ("png", "pdf"),
    dpi: int = 300,
) -> dict[str, Path]:
    """Export a figure plus an optional caption/context sidecar."""

    output_root = Path(output_dir)
    output_root.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for fmt in formats:
        paths[fmt] = save_figure(fig, output_root / f"{stem}.{fmt}", dpi=dpi)
    if context:
        context_path = output_root / f"{stem}.caption.md"
        context_path.write_text(context.as_markdown(), encoding="utf-8")
        paths["caption"] = context_path
    return paths


def export_word_figure(
    fig: plt.Figure,
    output_dir: str | Path,
    stem: str,
    *,
    context: FigureContext | None = None,
    spec: str | WordFigureSpec = "full_width",
) -> dict[str, Path]:
    """Export a Word-ready A4 PNG and context sidecar."""

    figure_spec = word_figure_spec(spec) if isinstance(spec, str) else spec
    if figure_spec.height_inches:
        fig.set_size_inches(figure_spec.width_inches, figure_spec.height_inches, forward=True)
    else:
        current = fig.get_size_inches()
        fig.set_size_inches(figure_spec.width_inches, current[1], forward=True)
    return export_figure_bundle(
        fig,
        output_dir,
        stem,
        context=context,
        formats=("png",),
        dpi=figure_spec.dpi,
    )


def _word_document_classes():
    """Return python-docx classes or raise a clear dependency error."""

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - dependency checked in setup
        raise RuntimeError("python-docx is required for Word export") from exc
    return Document, Inches


def _configure_a4_portrait(document: object) -> None:
    """Configure all sections as A4 portrait with the repo's Word margins."""

    _, inches = _word_document_classes()
    for section in document.sections:
        section.page_width = inches(A4_PORTRAIT_WIDTH_INCHES)
        section.page_height = inches(A4_PORTRAIT_HEIGHT_INCHES)
        section.left_margin = inches(A4_SIDE_MARGIN_INCHES)
        section.right_margin = inches(A4_SIDE_MARGIN_INCHES)
        section.top_margin = inches(A4_VERTICAL_MARGIN_INCHES)
        section.bottom_margin = inches(A4_VERTICAL_MARGIN_INCHES)


def _usable_page_width_inches(document: object) -> float:
    """Return the usable width for the current Word section."""

    section = document.sections[-1]
    usable_emu = section.page_width - section.left_margin - section.right_margin
    return max(float(usable_emu) / EMU_PER_INCH, 0.0)


def _set_cell_shading(cell: object, fill: str) -> None:
    """Set table-cell shading using python-docx XML primitives."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_borders(cell: object, color: str = "D8DDE6") -> None:
    """Set subtle borders around a table cell."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _add_caption_box(
    document: object,
    context: FigureContext,
    *,
    figure_number: int | None,
) -> None:
    """Add a contained caption box below a figure."""

    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = int(_usable_page_width_inches(document) * EMU_PER_INCH)
    cell = table.cell(0, 0)
    cell.width = int(_usable_page_width_inches(document) * EMU_PER_INCH)
    _set_cell_shading(cell, "F8FAFC")
    _set_cell_borders(cell)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05

    caption = context.caption_text(figure_number)
    prefix = "Figure" if figure_number is None else f"Figure {figure_number}"
    prefix_text = f"{prefix}."
    if caption.startswith(prefix_text):
        prefix_run = paragraph.add_run(prefix_text)
        prefix_run.bold = True
        paragraph.add_run(caption[len(prefix_text) :])
    else:
        paragraph.add_run(caption)

    document.add_paragraph()


def _add_figure_to_document(
    document: object,
    image_path: str | Path,
    *,
    context: FigureContext | None,
    spec: str | WordFigureSpec,
    figure_number: int | None,
) -> None:
    """Add one image and its context to a python-docx document."""

    _, inches = _word_document_classes()
    figure_spec = word_figure_spec(spec) if isinstance(spec, str) else spec
    width_inches = min(figure_spec.width_inches, _usable_page_width_inches(document))
    document.add_picture(str(image_path), width=inches(width_inches))
    if context:
        _add_caption_box(document, context, figure_number=figure_number)


def insert_figure_docx(
    image_path: str | Path,
    docx_path: str | Path,
    *,
    context: FigureContext | None = None,
    spec: str | WordFigureSpec = "full_width",
) -> Path:
    """Create a simple Word document containing one figure and caption."""

    document_class, _ = _word_document_classes()
    output_path = Path(docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = document_class()
    _configure_a4_portrait(document)
    _add_figure_to_document(
        document,
        image_path,
        context=context,
        spec=spec,
        figure_number=1 if context else None,
    )
    document.save(output_path)
    return output_path


def insert_figures_docx(
    entries: Sequence[WordFigureEntry],
    docx_path: str | Path,
    *,
    title: str = "Figure Proof Pack",
    page_breaks: bool = True,
) -> Path:
    """Create one Word document containing multiple figures and contexts."""

    if not entries:
        raise ValueError("at least one WordFigureEntry is required")

    document_class, _ = _word_document_classes()
    output_path = Path(docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = document_class()
    _configure_a4_portrait(document)
    if title:
        document.add_heading(title, level=0)
    for index, entry in enumerate(entries):
        if index and page_breaks:
            document.add_page_break()
        _add_figure_to_document(
            document,
            entry.image_path,
            context=entry.context,
            spec=entry.spec,
            figure_number=index + 1 if entry.context else None,
        )
    document.save(output_path)
    return output_path
